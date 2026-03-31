#!/usr/bin/env python3
"""
MedAgent Hub — 文档导出脚本
用法：
  python3 export_doc.py pdf  <markdown_text_base64> <output_path>
  python3 export_doc.py word <markdown_text_base64> <output_path>
"""
import sys, os, base64, re, tempfile

def md_to_html(md_text):
    """将 Markdown 转为带样式的 HTML"""
    try:
        import markdown
        body = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
    except ImportError:
        # 简单替换
        body = md_text.replace('\n', '<br>')
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<style>
  @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@400;500;700&display=swap');
  body {{
    font-family: 'Noto Sans SC', 'PingFang SC', 'Microsoft YaHei', sans-serif;
    font-size: 14px;
    line-height: 1.8;
    color: #1a1a2e;
    max-width: 800px;
    margin: 0 auto;
    padding: 40px 48px;
    background: #fff;
  }}
  h1 {{ font-size: 1.8em; font-weight: 700; color: #1a1a2e; border-bottom: 2px solid #e8e0f0; padding-bottom: 8px; margin-top: 0; }}
  h2 {{ font-size: 1.4em; font-weight: 600; color: #2d2d4e; margin-top: 1.5em; }}
  h3 {{ font-size: 1.15em; font-weight: 600; color: #3d3d5e; margin-top: 1.2em; }}
  p {{ margin: 0.6em 0; }}
  ul, ol {{ padding-left: 1.5em; }}
  li {{ margin: 0.3em 0; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
  th {{ background: #f0ebf8; color: #2d2d4e; font-weight: 600; padding: 8px 12px; border: 1px solid #d8d0e8; text-align: left; }}
  td {{ padding: 7px 12px; border: 1px solid #e8e0f0; }}
  tr:nth-child(even) td {{ background: #faf8fd; }}
  code {{ background: #f5f3ff; color: #7c3aed; padding: 2px 5px; border-radius: 3px; font-size: 0.88em; font-family: 'Courier New', monospace; }}
  pre {{ background: #1e1e2e; color: #cdd6f4; padding: 16px; border-radius: 8px; overflow-x: auto; }}
  pre code {{ background: none; color: inherit; padding: 0; }}
  blockquote {{ border-left: 3px solid #7c3aed; margin: 1em 0; padding: 0.5em 1em; background: #faf8fd; color: #555; }}
  strong {{ color: #2d2d4e; }}
  .footer {{ margin-top: 3em; padding-top: 1em; border-top: 1px solid #e8e0f0; font-size: 0.8em; color: #999; text-align: center; }}
</style>
</head>
<body>
{body}
<div class="footer">由 MedAgent Hub 生成 · medagenthub.com</div>
</body>
</html>"""

def export_pdf(md_text, output_path):
    html = md_to_html(md_text)
    with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
        f.write(html)
        tmp_html = f.name
    try:
        from weasyprint import HTML, CSS
        HTML(filename=tmp_html).write_pdf(output_path)
        print(f'OK:{output_path}')
    except Exception as e:
        print(f'ERROR:{e}', file=sys.stderr)
        sys.exit(1)
    finally:
        os.unlink(tmp_html)

def export_word(md_text, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()
        # 设置页面边距
        section = doc.sections[0]
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(11)

        lines = md_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            # 标题
            if line.startswith('### '):
                p = doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('# '):
                p = doc.add_heading(line[2:].strip(), level=1)
            # 表格（检测 | 开头的行）
            elif line.strip().startswith('|') and '|' in line[1:]:
                # 收集表格行
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                # 过滤分隔行
                data_lines = [l for l in table_lines if not re.match(r'^\s*\|[-: |]+\|\s*$', l)]
                if data_lines:
                    # 解析列
                    cols = [c.strip() for c in data_lines[0].split('|') if c.strip()]
                    table = doc.add_table(rows=1, cols=len(cols))
                    table.style = 'Table Grid'
                    # 表头
                    hdr = table.rows[0].cells
                    for j, col in enumerate(cols):
                        hdr[j].text = col
                        hdr[j].paragraphs[0].runs[0].bold = True if hdr[j].paragraphs[0].runs else False
                    # 数据行
                    for dl in data_lines[1:]:
                        row_cells = table.add_row().cells
                        vals = [c.strip() for c in dl.split('|') if c.strip()]
                        for j, val in enumerate(vals[:len(cols)]):
                            row_cells[j].text = val
                doc.add_paragraph('')
                continue
            # 列表项
            elif re.match(r'^[-*] ', line):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(re.sub(r'^\d+\. ', '', line).strip(), style='List Number')
            # 空行
            elif line.strip() == '':
                if i > 0 and lines[i-1].strip() != '':
                    doc.add_paragraph('')
            # 普通段落
            else:
                # 处理行内加粗 **text**
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.+?)\*\*', line)
                for j, part in enumerate(parts):
                    run = p.add_run(part)
                    if j % 2 == 1:  # 奇数位是加粗内容
                        run.bold = True
            i += 1

        # 页脚
        doc.add_paragraph('')
        footer_p = doc.add_paragraph('由 MedAgent Hub 生成 · medagenthub.com')
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.runs[0].font.size = Pt(9)
        footer_p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.save(output_path)
        print(f'OK:{output_path}')
    except Exception as e:
        print(f'ERROR:{e}', file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print('Usage: export_doc.py [pdf|word] <base64_text> <output_path>', file=sys.stderr)
        sys.exit(1)

    fmt = sys.argv[1]
    md_text = base64.b64decode(sys.argv[2]).decode('utf-8')
    output_path = sys.argv[3]

    if fmt == 'pdf':
        export_pdf(md_text, output_path)
    elif fmt == 'word':
        export_word(md_text, output_path)
    else:
        print(f'Unknown format: {fmt}', file=sys.stderr)
        sys.exit(1)
